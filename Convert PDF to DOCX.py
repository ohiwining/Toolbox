from pdfminer.pdfparser import PDFParser,PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from docx import Document

#Lib needed: pdfminer3k

document = Document()

def parse():
    fn = open('source.pdf', 'rb')
    parser = PDFParser(fn)
    doc = PDFDocument()
    parser.set_document(doc)
    doc.set_parser(parser)
    doc.initialize("")#set password for doc document
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        resource = PDFResourceManager()
        laparams = LAParams()
        device = PDFPageAggregator(resource, laparams=laparams)
        interpreter = PDFPageInterpreter(resource, device)
        for page in doc.get_pages():
            interpreter.process_page(page)
            layout = device.get_result()
            for out in layout:
                if hasattr(out, "get_text"):
                    content = out.get_text().replace(u'\xa0', u' ')
                    document.add_paragraph(
                        content, style='ListBullet'
                    )
                document.save('result.docx')


if __name__ == '__main__':
    parse()
