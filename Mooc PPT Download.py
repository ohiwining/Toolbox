import requests
import smtplib
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# 配置文件
PAGES = 25
Sender = "xxxx@163.com"  # 发邮件的地址
passwd = "xxxx"  # 发送者邮件的授权密码
SendTo = "xxxx@qq.com"  # 目标邮箱
FirstPngUrl = '''https://...ananas.chaoxing.com/.../1.png'''  # 网页上第一个ppt图片的链接


def Crawler():
    for page in range(1, PAGES + 1):
        url = FirstPngUrl.replace('1.png', '{}.png'.format(page))
        r = requests.get(url)
        with open('./PPTPIC/{}.png'.format(page), 'wb') as f:
            f.write(r.content)


def GenerateDocx():
    doc = docx.Document()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 图片居中设置
    run = paragraph.add_run("")
    for page in range(1, PAGES + 1):
        run.add_picture('./PPTPIC/{}.png'.format(page), width=Inches(6))
    doc.save("PPT.docx")


def SendMail():
    SMTPServer = "smtp.163.com"  # 模拟服务器  SMTP服务器
    message = MIMEMultipart()  # 创建一个带附件的实例
    message['From'] = Sender
    message['To'] = SendTo
    subject = 'PPT'
    message['Subject'] = subject
    message.attach(MIMEText('邮件发送测试', 'plain', 'utf-8'))  # 邮件正文内容
    att1 = MIMEText(open('PPT.docx', 'rb').read(), 'base64', 'utf-8')  # 构造附件1
    att1["Content-Type"] = 'application/octet-stream'
    att1["Content-Disposition"] = 'attachment; filename="a.docx"'  # 这里的filename可以任意写，写什么名字，邮件中显示什么名字
    message.attach(att1)
    mailServer = smtplib.SMTP(SMTPServer, 25)  # 25为端口号(邮件），0-1024都被系统占用了
    mailServer.login(Sender, passwd)  # 登录邮箱 需要的是，邮箱的地址和授权密码
    mailServer.sendmail(Sender, [SendTo], message.as_string())  # 发送文件
    print("Successful!")


def main():
    try:
        Crawler()
        print("爬取图片成功,图片保存到PPTPIC文件夹中")
    except:
        print("爬取图片失败")
        return
    try:
        GenerateDocx()
        print("生成doc文件成功,文件保存在源码根目录")
    except:
        print("生成doc文件失败")
        return
    try:
        SendMail()
        print("发送邮件成功,请查收")
    except:
        print("发送邮件失败")
        return


if __name__ == '__main__':
    main()
    print("End.")
