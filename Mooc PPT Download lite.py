import requests
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

# 配置文件
FirstPngUrl = '''https://...ananas.chaoxing.com/.../1.png'''  # 网页上第一个ppt图片的链接
PAGES = 25


def Crawler():
    for page in range(1, PAGES + 1):
        url = FirstPngUrl.replace('1.png', '{}.png'.format(page))
        r = requests.get(url)
        with open('./PPTPIC/{}.png'.format(page), 'wb') as f:
            f.write(r.content)


def GenerateDocx():
    doc = docx.Document()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 图片居中设置
    run = paragraph.add_run("")
    for page in range(1, PAGES + 1):
        run.add_picture('./PPTPIC/{}.png'.format(page), width=Inches(6))
    doc.save("PPT.docx")


def main():
    try:
        Crawler()
        print("爬取图片成功,图片保存到PPTPIC文件夹中")
    except:
        print("爬取图片失败")
        return
    try:
        GenerateDocx()
        print("生成doc文件成功,文件保存在源码根目录")
    except:
        print("生成doc文件失败")
        return


if __name__ == '__main__':
    main()
    print("End.")
